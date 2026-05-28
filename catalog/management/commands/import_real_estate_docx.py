"""
Импорт текстов договоров из DOCX в опубликованные шаблоны каталога «Недвижимость».

Ожидаемая структура: каталог ``documents/Шаблоны`` (в корне проекта), вложенные
подпапки допускаются. Для одного шаблона — основной файл + файлы с
«приложение N» в имени; порядок: основной договор, затем приложения по номеру.

Запуск (из корня проекта, с активированным venv)::

    python manage.py import_real_estate_docx
    python manage.py import_real_estate_docx --dry-run

Не удаляет и не переименовывает ``DocumentType`` / ``Template``; обновляет только
поле ``body`` у существующей опубликованной ``TemplateVersion`` с максимальным
``version_number`` для каждого сопоставленного типа документа.
"""

from __future__ import annotations

import re
from collections import defaultdict
from difflib import SequenceMatcher
from pathlib import Path

from django.conf import settings
from django.core.management.base import BaseCommand
from django.db import transaction

from catalog.models import DocumentType, TemplateVersion


def _norm(s: str) -> str:
    s = (s or "").lower().replace("ё", "е")
    s = re.sub(r"[^\w\s]+", " ", s, flags=re.UNICODE)
    return " ".join(s.split())


def _docx_to_plain(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    lines.append(t)
    return "\n".join(lines)


def _base_stem(stem: str) -> str:
    """Имя файла без расширения → ключ группы (без блока «приложение …»)."""
    s = re.sub(r"^\d+\.\s*", "", stem)
    parts = re.split(r"(?i)\s*приложени[ея]", s, maxsplit=1)
    return parts[0].strip(" -._")


def _appendix_sort_key(stem: str) -> tuple[int, int, str]:
    m = re.search(r"(?i)приложени[ея]\s*(\d+)", stem)
    if m:
        return (1, int(m.group(1)), stem.lower())
    return (0, 0, stem.lower())


def _best_document_type(base: str, candidates: list[DocumentType]) -> DocumentType | None:
    nb = _norm(base)
    if not nb:
        return None
    best: tuple[float, DocumentType] | None = None
    for dt in candidates:
        nn = _norm(dt.name)
        r = SequenceMatcher(None, nb, nn).ratio()
        if nb in nn or nn in nb:
            r = max(r, 0.82)
        if best is None or r > best[0]:
            best = (r, dt)
    if best is None or best[0] < 0.38:
        return None
    return best[1]


class Command(BaseCommand):
    help = "Импорт DOCX из documents/Шаблоны в тела шаблонов недвижимости"

    def add_arguments(self, parser):
        parser.add_argument(
            "--root",
            type=str,
            default="",
            help="Каталог с DOCX (по умолчанию: BASE_DIR/documents/Шаблоны)",
        )
        parser.add_argument(
            "--dry-run",
            action="store_true",
            help="Только показать сопоставления, без записи в БД",
        )

    def handle(self, *args, **options):
        root = Path(options["root"] or (settings.BASE_DIR / "documents" / "Шаблоны"))
        dry = options["dry_run"]

        if not root.is_dir():
            self.stderr.write(self.style.WARNING(f"Каталог не найден: {root}"))
            return

        docx_files = sorted(root.rglob("*.docx"))
        if not docx_files:
            self.stderr.write(self.style.WARNING(f"Нет файлов .docx в {root}"))
            return

        real_estate_types = list(
            DocumentType.objects.filter(domain__slug="real-estate", is_active=True).order_by(
                "sort_order", "name"
            )
        )
        if not real_estate_types:
            self.stderr.write(self.style.ERROR("В БД нет типов документов для домена real-estate"))
            return

        groups: dict[str, list[Path]] = defaultdict(list)
        for p in docx_files:
            key = _base_stem(p.stem)
            groups[key].append(p)

        for key, paths in groups.items():
            paths.sort(key=lambda x: _appendix_sort_key(x.stem))

        matched = 0
        skipped = 0
        to_write: list[tuple[DocumentType, str, list[Path]]] = []

        for base_key, paths in sorted(groups.items(), key=lambda kv: kv[0].lower()):
            dt = _best_document_type(base_key, real_estate_types)
            if dt is None:
                self.stdout.write(f"[пропуск] нет сопоставления для группы «{base_key}» ({len(paths)} файл.)")
                skipped += 1
                continue
            chunks: list[str] = []
            for p in paths:
                try:
                    chunks.append(_docx_to_plain(p))
                except Exception as exc:
                    self.stderr.write(self.style.ERROR(f"Ошибка чтения {p}: {exc}"))
                    continue
            if not chunks:
                self.stdout.write(f"[пропуск] не удалось прочитать ни одного DOCX в группе «{base_key}»")
                skipped += 1
                continue
            body = "\n\n".join(c for c in chunks if c.strip())
            to_write.append((dt, body, paths))
            matched += 1

        self.stdout.write(self.style.NOTICE(f"Сопоставлено групп: {matched}, без пары: {skipped}"))

        if dry:
            for dt, body, paths in to_write:
                self.stdout.write(f"  {dt.name} <- {len(paths)} file(s), text length {len(body)}")
            return

        with transaction.atomic():
            for dt, body, paths in to_write:
                tv = (
                    TemplateVersion.objects.filter(
                        template__document_type=dt,
                        template__is_active=True,
                        is_published=True,
                    )
                    .select_related("template")
                    .order_by("-version_number")
                    .first()
                )
                if tv is None:
                    self.stderr.write(self.style.WARNING(f"Нет опубликованной версии шаблона для: {dt.name}"))
                    continue
                tv.body = body
                tv.save(update_fields=["body"])
                self.stdout.write(self.style.SUCCESS(f"Обновлён шаблон: {dt.name} ({len(paths)} файл.)"))
