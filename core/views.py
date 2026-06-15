import json
from datetime import date
from io import BytesIO

from django.conf import settings as django_settings
from django.contrib.auth.decorators import login_required
from django.core.cache import cache
from django.http import Http404, HttpResponse, JsonResponse
from django.views.decorators.csrf import ensure_csrf_cookie
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.utils.text import slugify
from django.views.decorators.http import require_GET, require_POST

from docx import Document

from catalog.models import DocumentType, LegalDomain, TemplateVersion
from documents.models import DocumentFieldValue, UserDocument

from .legal_docs import paragraphs_to_safe_html, read_docx_plain_paragraphs, resolve_legal_doc_path
from .services import (
    call_openai_chat,
    extract_variable_keys,
    is_profanity_or_gibberish,
    merge_template_body,
    search_document_types,
)

REAL_ESTATE_GROUPS = [
    {
        "name": "Договоры по отчуждению жилых помещений",
        "documents": [
            "Договор купли-продажи жилого помещения",
            "Договор мены жилых помещений",
            "Договор дарения жилого помещения",
            "Договор ренты (постоянной, пожизненной)",
            "Договор пожизненного содержания с иждивением",
            "Договор уступки прав требования (цессии) по ДДУ",
            "Договор купли-продажи жилого помещения с рассрочкой платежа",
        ],
    },
    {
        "name": "Договоры найма и пользования жилыми помещениями",
        "documents": [
            "Договор социального найма жилого помещения",
            "Договор найма жилого помещения (коммерческий наем)",
            "Договор найма специализированного жилого помещения (служебное жилье, общежитие и др.)",
            "Договор поднайма жилого помещения",
            "Договор безвозмездного пользования жилым помещением (ссуды)",
            "Договор аренды жилого помещения (для юридических лиц — только для проживания граждан)",
            "Договор краткосрочного найма (до 1 года)",
        ],
    },
    {
        "name": "Иные договоры с жилыми помещениями",
        "documents": [
            "Договор доверительного управления жилым помещением",
            "Договор ипотеки (залога) жилого помещения",
            "Договор пожизненного проживания",
            "Договор о предоставлении жилого помещения в пользование членам семьи собственника",
            "Договор об определении порядка пользования жилым помещением (между сособственниками)",
            "Соглашение о разделе (выделе долей) жилого помещения",
            "Договор обмена жилыми помещениями (между нанимателями по договору соцнайма)",
            "Договор ренты с условием пожизненного проживания",
            "Договор купли-продажи жилого помещения с условием пожизненного проживания продавца",
        ],
    },
]

STUB_DOMAINS = [
    {"name": "Семейное право", "description": "Раздел в разработке для следующих версий сервиса."},
    {"name": "Трудовое право", "description": "Раздел в разработке для следующих версий сервиса."},
    {"name": "Корпоративное право", "description": "Раздел в разработке для следующих версий сервиса."},
]


def _normalize_name(value):
    return " ".join(value.lower().split())


VARIABLE_LABELS = {
    "party_one": "Сторона 1",
    "party_two": "Сторона 2",
    "subject": "Предмет",
    "date": "Дата",
}


def _humanize_variable_key(key: str) -> str:
    normalized = str(key).strip().lower()
    if normalized in VARIABLE_LABELS:
        return VARIABLE_LABELS[normalized]
    return str(key).replace("_", " ").strip().capitalize()


def _build_catalog_tree():
    domains = list(
        LegalDomain.objects.filter(is_active=True, slug__in=["real-estate", "construction"])
        .prefetch_related("document_types")
        .order_by("sort_order", "name")
    )
    docs_by_domain_slug = {}
    for domain in domains:
        docs_by_domain_slug[domain.slug] = {
            _normalize_name(doc.name): doc for doc in domain.document_types.all()
        }

    # ID документов у которых есть размеченный .docx шаблон
    from catalog.models import Template as CatalogTemplate
    docx_ids = set(
        TemplateVersion.objects
        .filter(is_published=True, template__is_active=True)
        .exclude(docx_file="")
        .exclude(docx_file=None)
        .values_list("template__document_type_id", flat=True)
    )

    catalog_tree = []
    for domain in domains:
        groups = []
        if domain.slug == "real-estate":
            domain_docs = docs_by_domain_slug.get(domain.slug, {})
            for group in REAL_ESTATE_GROUPS:
                items = []
                for doc_name in group["documents"]:
                    doc_type = domain_docs.get(_normalize_name(doc_name))
                    items.append(
                        {
                            "name": doc_name,
                            "doc_type": doc_type,
                            "has_fill_form": doc_type is not None and doc_type.id in docx_ids,
                        }
                    )
                groups.append({"name": group["name"], "items": items})
        else:
            items = [
                {
                    "name": doc.name,
                    "doc_type": doc,
                    "has_fill_form": doc.id in docx_ids,
                }
                for doc in domain.document_types.all()
                if doc.is_active
            ]
            groups.append({"name": "Документы по строительству", "items": items})

        catalog_tree.append({"domain": domain, "groups": groups})

    return catalog_tree


@ensure_csrf_cookie
def home(request):
    featured = list(
        DocumentType.objects.filter(domain__slug="real-estate", is_active=True).order_by("sort_order")[:6]
    )
    return render(request, "core/home.html", {"featured_templates": featured})


@ensure_csrf_cookie
@login_required
def cabinet(request):
    catalog_tree = _build_catalog_tree()
    mvp_domains = list(
        LegalDomain.objects.filter(is_active=True, slug__in=["real-estate", "construction"]).order_by(
            "sort_order", "name"
        )
    )
    selected_slug = request.GET.get("domain") or (mvp_domains[0].slug if mvp_domains else "real-estate")
    selected_domain = next((d for d in mvp_domains if d.slug == selected_slug), None)
    if selected_domain is None and mvp_domains:
        selected_domain = mvp_domains[0]
        selected_slug = selected_domain.slug

    document_cards = []
    if selected_domain:
        document_cards = list(
            DocumentType.objects.filter(domain=selected_domain, is_active=True).order_by(
                "sort_order", "name"
            )
        )

    workspace = None
    doc_id = request.GET.get("doc")
    if doc_id:
        ud = (
            UserDocument.objects.filter(id=doc_id, user=request.user)
            .select_related("template_version", "template_version__template", "template_version__template__document_type")
            .first()
        )
        if ud:
            body = ud.template_version.body
            keys = extract_variable_keys(body)
            values = {fv.variable_key: fv.value for fv in ud.field_values.all()}
            merged = merge_template_body(body, {k: values.get(k, "") for k in keys})
            workspace = {
                "document": ud,
                "template_body": body,
                "variable_keys": keys,
                "variable_rows": [
                    {"key": k, "label": _humanize_variable_key(k), "value": values.get(k, "")}
                    for k in keys
                ],
                "values": values,
                "merged_preview": merged,
                "template_version_label": f"Версия шаблона {ud.template_version.version_number}",
                "domain_slug": ud.template_version.template.document_type.domain.slug,
            }

    return render(
        request,
        "core/cabinet.html",
        {
            "catalog_tree": catalog_tree,
            "mvp_domains": mvp_domains,
            "stub_domains": STUB_DOMAINS,
            "selected_domain": selected_domain,
            "selected_domain_slug": selected_slug,
            "document_cards": document_cards,
            "workspace": workspace,
        },
    )


@login_required
def start_document(request, document_type_id):
    document_type = get_object_or_404(DocumentType, id=document_type_id, is_active=True)
    template_version = (
        TemplateVersion.objects.filter(
            template__document_type=document_type,
            template__is_active=True,
            is_published=True,
        )
        .select_related("template")
        .order_by("-version_number")
        .first()
    )
    if template_version is None:
        return redirect(f"{reverse('core:cabinet')}?domain={document_type.domain.slug}")

    keys = extract_variable_keys(template_version.body)
    if keys:
        ud = UserDocument.objects.create(
            user=request.user,
            template_version=template_version,
            title=document_type.name,
            content=template_version.body,
            workspace_phase=UserDocument.WorkspacePhase.VARIABLES,
        )
        DocumentFieldValue.objects.bulk_create(
            [
                DocumentFieldValue(
                    user_document=ud,
                    variable_key=key,
                    variable_label=key,
                    value="",
                    is_required=True,
                )
                for key in keys
            ]
        )
    else:
        merged = merge_template_body(template_version.body, {})
        ud = UserDocument.objects.create(
            user=request.user,
            template_version=template_version,
            title=document_type.name,
            content=merged,
            workspace_phase=UserDocument.WorkspacePhase.EDITING,
        )

    return redirect(f"{reverse('core:cabinet')}?doc={ud.id}&domain={document_type.domain.slug}")


@login_required
def download_template_docx(request, document_type_id):
    document_type = get_object_or_404(DocumentType, id=document_type_id, is_active=True)
    template_version = (
        TemplateVersion.objects.filter(
            template__document_type=document_type,
            template__is_active=True,
            is_published=True,
        )
        .order_by("-version_number")
        .first()
    )
    if template_version is None:
        raise Http404("Шаблон документа не найден")

    # Если есть размеченный .docx — перенаправляем на форму заполнения
    if template_version.docx_file:
        return redirect(
            reverse("core:fill_docx_form", args=[document_type_id])
        )

    # Иначе — старый путь: генерируем docx из plain text
    doc = Document()
    doc.add_heading(document_type.name, level=1)
    for line in template_version.body.splitlines():
        text = line.strip()
        if text:
            doc.add_paragraph(text)
        else:
            doc.add_paragraph("")

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"{slugify(document_type.name, allow_unicode=True)}.docx"
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _extract_docx_variables(path: str) -> list[str]:
    """Извлекает все {{ переменные }} из .docx файла."""
    import re
    doc = Document(path)
    all_text = " ".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += " ".join(p.text for p in cell.paragraphs)
    # Заголовки и подвалы
    for section in doc.sections:
        try:
            all_text += " ".join(p.text for p in section.header.paragraphs)
            all_text += " ".join(p.text for p in section.footer.paragraphs)
        except Exception:
            pass
    return sorted(set(re.findall(r"\{\{\s*(\w+)\s*\}\}", all_text)))


def _extract_docx_text_blocks(path: str):
    """
    Извлекает форматированные блоки текста из .docx в порядке документа.
    Каждый блок — dict с ключами:
      type: 'p' | 'tr'
      segments: [('text'|'var'|'blank', str), ...]
      css_class: str   (Tailwind классы)
      para_style: str  (inline CSS)
    """
    import re as _re
    from docx.text.paragraph import Paragraph as _DocxPara
    from docx.table import Table as _DocxTable
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    SEGMENT_RE = _re.compile(r'\{\{\s*(\w+)\s*\}\}|«_+»|_{2,}')

    def parse_segments(text):
        segs, last = [], 0
        for m in SEGMENT_RE.finditer(text):
            if m.start() > last:
                segs.append(('text', text[last:m.start()]))
            g = m.group()
            if g.startswith('{{'):
                segs.append(('var', m.group(1)))
            else:
                segs.append(('blank', g))
            last = m.end()
        if last < len(text):
            segs.append(('text', text[last:]))
        return segs or [('text', '')]

    doc = Document(path)
    blocks = []

    for child in doc.element.body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            para = _DocxPara(child, doc)
            text = para.text.strip()
            if not text:
                continue

            sname = (para.style.name or 'Normal') if para.style else 'Normal'
            align = para.alignment
            bold = any(r.bold for r in para.runs if r.text.strip())
            # Заголовок: по стилю ИЛИ короткий + центр + жирный
            heading = (
                'heading' in sname.lower()
                or 'заголовок' in sname.lower()
                or (bold and align == WD_ALIGN_PARAGRAPH.CENTER and len(text) < 120)
            )
            if heading:
                bold = True

            css = ['mb-1']
            if bold:
                css.append('font-bold')
            if align == WD_ALIGN_PARAGRAPH.CENTER:
                css.append('text-center')
            elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                css.append('text-right')
            elif align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                css.append('text-justify')

            # Отступ абзаца
            try:
                li = para.paragraph_format.left_indent
                indent_em = round(li / 360000 * 0.5, 2) if li and li > 0 else 0
            except Exception:
                indent_em = 0

            if indent_em > 0:
                pstyle = f'padding-left: {indent_em}em;'
            elif (
                not heading
                and len(text) > 60
                and align not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)
            ):
                pstyle = 'text-indent: 1.25em;'
            else:
                pstyle = ''

            blocks.append({
                'type': 'p',
                'segments': parse_segments(text),
                'css_class': ' '.join(css),
                'para_style': pstyle,
            })

        elif tag == 'tbl':
            tbl = _DocxTable(child, doc)
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                # Убираем дубликаты слитых ячеек
                deduped: list[str] = []
                for c in cells:
                    if not deduped or c != deduped[-1]:
                        deduped.append(c)
                if deduped:
                    blocks.append({
                        'type': 'tr',
                        'segments': parse_segments('   '.join(deduped)),
                        'css_class': 'mb-1',
                        'para_style': '',
                    })

    return blocks


def _collect_docx_variables(docx_path, attachments):
    """Все переменные основного документа + эксклюзивные переменные приложений.
    Возвращает (variables, var_to_att)."""
    variables = _extract_docx_variables(docx_path)
    var_to_att: dict[str, int] = {}
    vars_set = set(variables)
    for att in attachments:
        try:
            for v in _extract_docx_variables(att.docx_file.path):
                if v not in vars_set:
                    variables.append(v)
                    vars_set.add(v)
                    var_to_att[v] = att.id
        except Exception:
            pass
    return variables, var_to_att


def _render_filled_files(docx_path, variables, values, included_atts, slug_name):
    """Рендерит основной .docx, при наличии выбранных приложений — упаковывает в .zip.
    Возвращает (bytes, content_type, filename)."""
    import zipfile
    from docxtpl import DocxTemplate

    context = {v: values.get(v, "") for v in variables}
    if not included_atts:
        tpl = DocxTemplate(docx_path)
        tpl.render(context)
        buf = BytesIO()
        tpl.save(buf)
        return (
            buf.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{slug_name}.docx",
        )
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        tpl = DocxTemplate(docx_path)
        tpl.render(context)
        main_buf = BytesIO()
        tpl.save(main_buf)
        zf.writestr(f"{slug_name}.docx", main_buf.getvalue())
        for att in included_atts:
            try:
                att_tpl = DocxTemplate(att.docx_file.path)
                att_tpl.render(context)
                att_buf = BytesIO()
                att_tpl.save(att_buf)
                zf.writestr(f"{slugify(att.title, allow_unicode=True)}.docx", att_buf.getvalue())
            except Exception:
                pass
    return zip_buffer.getvalue(), "application/zip", f"{slug_name}.zip"


def _attachment_response(content, content_type, filename):
    response = HttpResponse(content, content_type=content_type)
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


@login_required
def fill_docx_form(request, document_type_id):
    """GET: форма заполнения (с ?doc=<id> подгружает сохранённые значения);
    POST: генерирует docx/zip через docxtpl."""
    from .variables import group_variables

    document_type = get_object_or_404(DocumentType, id=document_type_id, is_active=True)
    template_version = (
        TemplateVersion.objects.filter(
            template__document_type=document_type,
            template__is_active=True,
            is_published=True,
        )
        .order_by("-version_number")
        .first()
    )
    if template_version is None or not template_version.docx_file:
        raise Http404("Размеченный шаблон не найден")

    docx_path = template_version.docx_file.path
    attachments = list(template_version.attachments.all())
    variables, var_to_att = _collect_docx_variables(docx_path, attachments)

    if request.method == "POST":
        values = {v: request.POST.get(v, "") for v in variables}
        slug_name = slugify(document_type.name, allow_unicode=True)
        included_atts = [
            att for att in attachments
            if request.POST.get(f"include_att_{att.id}") == "on"
        ]
        content, ctype, filename = _render_filled_files(
            docx_path, variables, values, included_atts, slug_name
        )
        return _attachment_response(content, ctype, filename)

    # ── Редактирование сохранённого документа: подгружаем значения ──
    user_doc = None
    saved_values: dict[str, str] = {}
    saved_att_included: set[int] = set()
    doc_id = request.GET.get("doc")
    if doc_id:
        user_doc = UserDocument.objects.filter(id=doc_id, user=request.user).first()
        if user_doc:
            for fv in user_doc.field_values.all():
                if fv.variable_key.startswith("__include_att_"):
                    if fv.value == "on":
                        try:
                            saved_att_included.add(int(fv.variable_key[len("__include_att_"):]))
                        except ValueError:
                            pass
                else:
                    saved_values[fv.variable_key] = fv.value

    # GET: собираем поля формы с группировкой
    text_blocks = _extract_docx_text_blocks(docx_path)
    var_order: list[str] = []
    _seen: set[str] = set()
    for _block in text_blocks:
        for _kind, _content in _block['segments']:
            if _kind == "var" and _content not in _seen:
                var_order.append(_content)
                _seen.add(_content)

    attachment_previews = []
    for att in attachments:
        try:
            att_blocks = _extract_docx_text_blocks(att.docx_file.path)
        except Exception:
            att_blocks = []
        attachment_previews.append({"attachment": att, "blocks": att_blocks})

    field_groups = group_variables(variables, saved_values, var_order=var_order)
    for _sec, _fields in field_groups:
        for _field in _fields:
            _field['att_id'] = var_to_att.get(_field['name'])

    return render(request, "core/fill_docx_form.html", {
        "document_type": document_type,
        "field_groups": field_groups,
        "total_fields": len(variables),
        "text_blocks": text_blocks,
        "attachments": attachments,
        "attachment_previews": attachment_previews,
        "user_doc": user_doc,
        "saved_att_included": saved_att_included,
        "is_editing": user_doc is not None,
    })


@login_required
@require_POST
def save_filled_document(request, document_type_id):
    """Сохраняет введённые в форму значения как документ пользователя («Мои документы»)."""
    document_type = get_object_or_404(DocumentType, id=document_type_id, is_active=True)
    template_version = (
        TemplateVersion.objects.filter(
            template__document_type=document_type,
            template__is_active=True,
            is_published=True,
        )
        .order_by("-version_number")
        .first()
    )
    if template_version is None or not template_version.docx_file:
        raise Http404("Размеченный шаблон не найден")

    docx_path = template_version.docx_file.path
    attachments = list(template_version.attachments.all())
    variables, _ = _collect_docx_variables(docx_path, attachments)

    user_doc_id = request.POST.get("user_doc_id")
    ud = None
    if user_doc_id:
        ud = UserDocument.objects.filter(id=user_doc_id, user=request.user).first()
    if ud is None:
        ud = UserDocument.objects.create(
            user=request.user,
            template_version=template_version,
            title=document_type.name,
            content="",
            status=UserDocument.Status.DRAFT,
            workspace_phase=UserDocument.WorkspacePhase.VARIABLES,
        )

    ud.field_values.all().delete()
    rows = [
        DocumentFieldValue(
            user_document=ud, variable_key=k, variable_label=k,
            value=request.POST.get(k, ""), is_required=False,
        )
        for k in variables
    ]
    for att in attachments:
        rows.append(DocumentFieldValue(
            user_document=ud,
            variable_key=f"__include_att_{att.id}",
            variable_label="",
            value="on" if request.POST.get(f"include_att_{att.id}") == "on" else "",
            is_required=False,
        ))
    DocumentFieldValue.objects.bulk_create(rows)
    ud.last_modified_at = timezone.now()
    ud.save()

    return JsonResponse({"ok": True, "redirect": reverse("core:my_documents"), "doc_id": ud.id})


@login_required
def export_filled_document(request, doc_id):
    """Перескачивает сохранённый документ (docx/zip) из сохранённых значений формы."""
    ud = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    tv = ud.template_version
    if not tv.docx_file:
        # старый текстовый документ — экспорт как раньше (plain docx)
        return download_user_document_docx(request, doc_id)

    docx_path = tv.docx_file.path
    attachments = list(tv.attachments.all())
    variables, _ = _collect_docx_variables(docx_path, attachments)

    values: dict[str, str] = {}
    included_ids: set[int] = set()
    for fv in ud.field_values.all():
        if fv.variable_key.startswith("__include_att_"):
            if fv.value == "on":
                try:
                    included_ids.add(int(fv.variable_key[len("__include_att_"):]))
                except ValueError:
                    pass
        else:
            values[fv.variable_key] = fv.value
    included_atts = [att for att in attachments if att.id in included_ids]
    slug_name = slugify(ud.title, allow_unicode=True)
    content, ctype, filename = _render_filled_files(
        docx_path, variables, values, included_atts, slug_name
    )
    return _attachment_response(content, ctype, filename)


@login_required
def download_user_document_docx(request, doc_id):
    ud = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    doc = Document()
    doc.add_heading(ud.title, level=1)
    for line in (ud.content or "").splitlines():
        t = line.strip()
        doc.add_paragraph(t if t else "")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    filename = f"{slugify(ud.title, allow_unicode=True)}.docx"
    response = HttpResponse(
        output.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


def _parse_json(request):
    try:
        return json.loads(request.body.decode("utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return None


@login_required
@require_POST
def save_document_fields(request, doc_id):
    payload = _parse_json(request)
    if not isinstance(payload, dict):
        return JsonResponse({"ok": False, "error": "Некорректный JSON"}, status=400)

    ud = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    if ud.workspace_phase != UserDocument.WorkspacePhase.VARIABLES:
        return JsonResponse({"ok": False, "error": "Неверная фаза документа"}, status=400)

    values = payload.get("values")
    if not isinstance(values, dict):
        return JsonResponse({"ok": False, "error": "Ожидался объект values"}, status=400)

    keys = extract_variable_keys(ud.template_version.body)
    allowed = set(keys)
    for key, raw in values.items():
        if key not in allowed:
            continue
        val = "" if raw is None else str(raw)
        DocumentFieldValue.objects.update_or_create(
            user_document=ud,
            variable_key=key,
            defaults={"value": val, "variable_label": key, "is_required": True},
        )

    ud.last_modified_at = timezone.now()
    ud.save()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def finalize_variables(request, doc_id):
    ud = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    if ud.workspace_phase != UserDocument.WorkspacePhase.VARIABLES:
        return redirect(f"{reverse('core:cabinet')}?doc={ud.id}&domain={ud.template_version.template.document_type.domain.slug}")

    body = ud.template_version.body
    keys = extract_variable_keys(body)
    values = {fv.variable_key: fv.value for fv in ud.field_values.all()}
    missing = [k for k in keys if not (values.get(k) or "").strip()]
    if missing:
        return JsonResponse(
            {"ok": False, "error": "Заполните все обязательные поля", "missing": missing},
            status=400,
        )

    merged = merge_template_body(body, {k: (values.get(k) or "").strip() for k in keys})
    ud.content = merged
    ud.workspace_phase = UserDocument.WorkspacePhase.EDITING
    ud.last_modified_at = timezone.now()
    ud.save()
    return JsonResponse({"ok": True, "redirect": f"{reverse('core:cabinet')}?doc={ud.id}&domain={ud.template_version.template.document_type.domain.slug}"})


@login_required
@require_POST
def save_document_editor(request, doc_id):
    payload = _parse_json(request)
    if not isinstance(payload, dict):
        return JsonResponse({"ok": False, "error": "Некорректный JSON"}, status=400)

    ud = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    if ud.workspace_phase != UserDocument.WorkspacePhase.EDITING:
        return JsonResponse({"ok": False, "error": "Неверная фаза документа"}, status=400)

    title = payload.get("title")
    content = payload.get("content")
    if title is not None:
        ud.title = str(title)[:255]
    if content is not None:
        ud.content = str(content)
    ud.last_modified_at = timezone.now()
    ud.save()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def duplicate_document(request, doc_id):
    source = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    clone = UserDocument.objects.create(
        user=request.user,
        template_version=source.template_version,
        title=f"Копия — {source.title}"[:255],
        content=source.content,
        workspace_phase=source.workspace_phase,
        status=UserDocument.Status.DRAFT,
    )
    for fv in source.field_values.all():
        DocumentFieldValue.objects.create(
            user_document=clone,
            variable_key=fv.variable_key,
            variable_label=fv.variable_label,
            value=fv.value,
            is_required=fv.is_required,
        )
    next_url = request.POST.get("next")
    if next_url:
        return redirect(next_url)
    return redirect(f"{reverse('core:cabinet')}?doc={clone.id}&domain={source.template_version.template.document_type.domain.slug}")


@login_required
@require_POST
def delete_document(request, doc_id):
    ud = get_object_or_404(UserDocument, id=doc_id, user=request.user)
    domain_slug = ud.template_version.template.document_type.domain.slug
    next_url = request.POST.get("next")
    ud.delete()
    if next_url:
        return redirect(next_url)
    return redirect(f"{reverse('core:cabinet')}?domain={domain_slug}")


def get_ai_quota_state(request):
    limit_auth = int(getattr(django_settings, "AI_DAILY_LLM_LIMIT_AUTH", 15))
    limit_guest = int(getattr(django_settings, "AI_GUEST_LLM_LIMIT", 1))
    if request.user.is_authenticated:
        key = f"ai_llm_used:{request.user.id}:{date.today().isoformat()}"
        used = int(cache.get(key, 0) or 0)
        return {
            "limit": limit_auth,
            "used": used,
            "remaining": max(0, limit_auth - used),
        }
    used = int(request.session.get("ai_guest_llm_used", 0) or 0)
    return {
        "limit": limit_guest,
        "used": used,
        "remaining": max(0, limit_guest - used),
    }


def increment_ai_quota(request) -> None:
    if request.user.is_authenticated:
        key = f"ai_llm_used:{request.user.id}:{date.today().isoformat()}"
        if cache.add(key, 1, timeout=60 * 60 * 36):
            return
        try:
            cache.incr(key)
        except ValueError:
            cache.set(key, 1, timeout=60 * 60 * 36)
        return
    request.session["ai_guest_llm_used"] = int(request.session.get("ai_guest_llm_used", 0) or 0) + 1
    request.session.modified = True


def is_quota_question(text: str) -> bool:
    tl = (text or "").lower()
    if "сколько" in tl and ("запрос" in tl or "остал" in tl or "лимит" in tl):
        return True
    if "лимит" in tl and "?" in (text or ""):
        return True
    return False


def format_quota_answer(request):
    st = get_ai_quota_state(request)
    if request.user.is_authenticated:
        return (
            f"Сегодня у вас {st['remaining']} из {st['limit']} запросов к AI. "
            f"Использовано: {st['used']}."
        )
    return (
        f"Без входа в аккаунт доступно {st['limit']} запрос(ов). "
        f"Осталось: {st['remaining']}."
    )


def limit_exhausted_answer(request):
    if request.user.is_authenticated:
        return (
            "Вы использовали все запросы к AI на сегодня. "
            "Завтра лимит обновится. Полный доступ — в тарифах на главной странице (раздел «Тарифы»)."
        )
    return (
        "Бесплатный лимит запросов исчерпан. "
        "Чтобы продолжить работу с AI-помощником, оформите один из тарифов — раздел «Тарифы» на главной странице."
    )


@require_GET
def ai_quota(request):
    return JsonResponse(get_ai_quota_state(request))


@require_GET
def legal_document_preview(request, slug):
    if slug not in ("privacy", "consent"):
        return JsonResponse({"error": "not_found"}, status=404)
    titles = {
        "privacy": "Политика конфиденциальности",
        "consent": "Согласие на обработку персональных данных",
    }
    path = resolve_legal_doc_path(slug)
    if not path or not path.exists():
        return JsonResponse(
            {
                "title": titles[slug],
                "html": (
                    "<p>Файл документа не найден. Укажите путь в переменных окружения "
                    "<code>LEGAL_PRIVACY_PATH</code> / <code>LEGAL_CONSENT_PATH</code> "
                    "или поместите DOCX в папку <code>legal_docs</code> в корне проекта.</p>"
                ),
            }
        )
    try:
        paragraphs = read_docx_plain_paragraphs(path)
        html = paragraphs_to_safe_html(paragraphs)
    except Exception:
        return JsonResponse(
            {
                "title": titles[slug],
                "html": "<p>Не удалось открыть документ. Проверьте, что файл не повреждён.</p>",
            }
        )
    return JsonResponse({"title": titles[slug], "html": html})


@login_required
def my_documents(request):
    docs = (
        UserDocument.objects.filter(
            user=request.user,
            status__in=[UserDocument.Status.DRAFT, UserDocument.Status.READY],
        )
        .select_related("template_version__template__document_type__domain")
    )
    items = []
    for d in docs:
        dt = d.template_version.template.document_type
        is_filled = bool(d.template_version.docx_file)
        if is_filled:
            edit_url = f"{reverse('core:fill_docx_form', args=[dt.id])}?doc={d.id}"
        else:
            edit_url = f"{reverse('core:cabinet')}?doc={d.id}&domain={dt.domain.slug}"
        items.append({
            "doc": d,
            "is_filled": is_filled,
            "edit_url": edit_url,
            "export_url": reverse("core:export_filled_document", args=[d.id]),
            "duplicate_url": reverse("core:duplicate_document", args=[d.id]),
            "delete_url": reverse("core:delete_document", args=[d.id]),
        })
    return render(request, "core/my_documents.html", {"documents": items})


def _doc_action_url(document_type) -> str:
    """Ссылка на документ: форма заполнения (если шаблон размечен .docx) или старый старт."""
    has_docx = (
        TemplateVersion.objects.filter(
            template__document_type=document_type,
            template__is_active=True,
            is_published=True,
        )
        .exclude(docx_file="")
        .exclude(docx_file=None)
        .exists()
    )
    if has_docx:
        return reverse("core:fill_docx_form", args=[document_type.id])
    return reverse("core:start_document", args=[document_type.id])


def ai_search_documents(request):
    if request.method != "POST":
        return JsonResponse({"error": "Method not allowed"}, status=405)
    payload = _parse_json(request)
    if not isinstance(payload, dict):
        return JsonResponse({"error": "Некорректный JSON"}, status=400)
    query = str(payload.get("query", "")).strip()
    template_results = search_document_types(query, limit=3)
    templates = [
        {
            "id": d.id,
            "name": d.name,
            "description": d.short_description or "",
            "start_url": _doc_action_url(d),
        }
        for d in template_results
    ]

    my_documents = []
    if request.user.is_authenticated:
        docs = (
            UserDocument.objects.filter(
                user=request.user,
                status__in=[UserDocument.Status.DRAFT, UserDocument.Status.READY],
                title__icontains=query,
            )
            .select_related("template_version__template__document_type__domain")
            .order_by("-updated_at")[:5]
        )
        my_documents = [
            {
                "id": doc.id,
                "name": doc.title,
                "updated_at": doc.updated_at.strftime("%d.%m.%Y, %H:%M"),
                "open_url": f"{reverse('core:cabinet')}?doc={doc.id}&domain={doc.template_version.template.document_type.domain.slug}",
            }
            for doc in docs
        ]

    return JsonResponse({"templates": templates, "my_documents": my_documents})


def ai_ask_llm(request):
    if request.method != "POST":
        return JsonResponse({"error": "Method not allowed"}, status=405)
    payload = _parse_json(request)
    if not isinstance(payload, dict):
        return JsonResponse({"error": "Некорректный JSON"}, status=400)
    question = str(payload.get("question", "")).strip()
    if not question:
        return JsonResponse({"error": "Пустой вопрос"}, status=400)

    quota = get_ai_quota_state(request)

    if is_quota_question(question):
        return JsonResponse(
            {
                "answer": format_quota_answer(request),
                "quota": get_ai_quota_state(request),
                "suggested_docs": [],
            }
        )

    if is_profanity_or_gibberish(question):
        return JsonResponse(
            {
                "answer": "Не удалось распознать запрос. Сформулируйте его по теме недвижимости или строительства.",
                "quota": quota,
                "suggested_docs": [],
            }
        )

    if quota["remaining"] <= 0:
        return JsonResponse(
            {
                "answer": limit_exhausted_answer(request),
                "quota": quota,
                "limited": True,
                "suggested_docs": [],
            }
        )

    # Ищем релевантные шаблоны: min_score=8 позволяет keyword-матчам работать.
    # Синонимы не хардкодятся — добавляются в поле keywords документа через админку.
    relevant_docs = search_document_types(question, limit=3, min_score=8)

    # Поиск в документах пользователя
    user_docs = []
    if request.user.is_authenticated:
        from documents.models import UserDocument
        qs = (
            UserDocument.objects.filter(
                user=request.user,
                status__in=["draft", "ready"],
                title__icontains=question,
            )
            .select_related("template_version__template__document_type__domain")
            .order_by("-updated_at")[:3]
        )
        user_docs = [
            {
                "name": d.title,
                "updated_at": d.updated_at.strftime("%d.%m.%Y"),
                "open_url": f"{reverse('core:cabinet')}?doc={d.id}&domain={d.template_version.template.document_type.domain.slug}",
            }
            for d in qs
        ]

    try:
        answer = call_openai_chat(question, context_docs=relevant_docs or None)
    except TimeoutError:
        fallback_docs = [
            {"name": d.name, "start_url": _doc_action_url(d)}
            for d in relevant_docs
        ]
        fallback_answer = (
            "AI-сервис сейчас недоступен, но в каталоге нашлись подходящие шаблоны."
            if relevant_docs else
            "В данный момент сервис недоступен. Попробуйте позже или воспользуйтесь подбором шаблона в разделе «Шаблоны»."
        )
        return JsonResponse(
            {
                "answer": fallback_answer,
                "quota": quota,
                "suggested_docs": fallback_docs,
                "my_documents": user_docs,
            }
        )

    increment_ai_quota(request)
    quota_after = get_ai_quota_state(request)

    suggested_docs = [
        {
            "name": d.name,
            "start_url": _doc_action_url(d),
        }
        for d in relevant_docs
    ]

    payload_out = {
        "answer": answer,
        "quota": quota_after,
        "suggested_docs": suggested_docs,
        "my_documents": user_docs,
    }
    if 0 < quota_after["remaining"] <= 5:
        payload_out["quota_warning"] = (
            f"Осталось {quota_after['remaining']} запросов к AI до конца дня."
        )
    return JsonResponse(payload_out)
